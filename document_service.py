import asyncio
from pathlib import Path
from typing import List
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
import re
import config
from models import OfertaRequest


class DocumentService:
    """Serwis do generowania dokumentów DOCX z szablonów"""

    def __init__(self):
        self.template_path = config.TEMPLATE_FILE
        self.produkty_dir = config.PRODUKTY_DIR

    async def generate_offer(self, request: OfertaRequest, output_path: Path) -> Path:
        """
        Generuje ofertę DOCX na podstawie szablonu i danych z request

        Args:
            request: Dane do wypełnienia oferty
            output_path: Ścieżka do zapisu wygenerowanego pliku

        Returns:
            Path: Ścieżka do wygenerowanego pliku DOCX
        """
        # Wczytaj szablon
        doc = Document(str(self.template_path))

        # Zamień placeholdery w dokumencie
        await self._replace_placeholders(doc, request)

        # Znajdź punkt wstrzyknięcia produktów (po paragrafie z opisem)
        injection_index = await self._find_injection_point(doc, request)

        # Wstaw produkty jako osobne dokumenty
        if injection_index is not None and request.produkty:
            await self._inject_products(doc, injection_index, request.produkty)

        # Zapisz dokument
        doc.save(str(output_path))
        return output_path

    async def _replace_placeholders(self, doc: Document, request: OfertaRequest):
        """Zamienia placeholdery w dokumencie na wartości z request"""

        # Mapa placeholderów -> wartości
        replacements = {
            "{{KLIENT(NIP)}}": request.nip,
            "{{Oferta z dnia}}": request.data_oferty,
            "{{waznado}}": request.wazna_do,
            "{{firmaM}}": request.firma,
            "{{temat}}": request.temat or "",
            "{{kategoria}}": request.kategoria or "",
            "{{opis}}": request.opis or "",
            "{{cena}}": f"{request.cena:.2f}" if request.cena else "",
            "{{RBG}}": str(request.rbg) if request.rbg else "",
            "{{uzasadnienie}}": request.uzasadnienie or "",
        }

        # Zamień w paragrafach
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    await self._replace_in_paragraph(paragraph, key, value)

        # Zamień w tabelach
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                await self._replace_in_paragraph(paragraph, key, value)

    async def _replace_in_paragraph(self, paragraph, search: str, replace: str):
        """Zamienia tekst w paragrafie zachowując formatowanie"""
        if search in paragraph.text:
            # Prostsze podejście: zastępujemy cały tekst
            inline = paragraph.runs
            for run in inline:
                if search in run.text:
                    run.text = run.text.replace(search, replace)

    async def _find_injection_point(self, doc: Document, request: OfertaRequest) -> int:
        """
        Znajduje indeks paragrafu, po którym należy wstrzyknąć produkty
        Szuka paragrafu zawierającego "Opis: \t\t{{opis}}" lub już wypełniony opis
        """
        marker_pattern = re.compile(r"Opis:\s+")

        for idx, paragraph in enumerate(doc.paragraphs):
            if marker_pattern.search(paragraph.text):
                # Znaleziono punkt wstrzyknięcia - zwróć indeks NASTĘPNEGO paragrafu
                return idx + 1

        # Jeśli nie znaleziono, wstaw po 10 paragrafie (zabezpieczenie)
        return 10

    async def _inject_products(self, doc: Document, insertion_index: int, product_files: List[str]):
        """
        Wstrzykuje produkty jako osobne dokumenty w określonym miejscu

        Args:
            doc: Główny dokument
            insertion_index: Indeks po którym wstrzyknąć produkty
            product_files: Lista nazw plików produktów
        """
        # Pobierz element XML głównego dokumentu
        main_body = doc._element.body

        # Dla każdego produktu
        for product_file in product_files:
            product_path = self.produkty_dir / product_file

            if not product_path.exists():
                print(f"⚠️ Ostrzeżenie: Produkt {product_file} nie istnieje, pomijam")
                continue

            # Wczytaj dokument produktu
            product_doc = Document(str(product_path))

            # Dodaj page break przed produktem
            page_break = parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:br w:type="page"/></w:r></w:p>')
            main_body.insert(insertion_index, page_break)
            insertion_index += 1

            # Wstaw wszystkie elementy z dokumentu produktu
            for element in product_doc._element.body:
                # Klonuj element i wstaw do głównego dokumentu
                cloned_element = element
                main_body.insert(insertion_index, cloned_element)
                insertion_index += 1

        # Dodaj page break po ostatnim produkcie
        page_break = parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:br w:type="page"/></w:r></w:p>')
        main_body.insert(insertion_index, page_break)
