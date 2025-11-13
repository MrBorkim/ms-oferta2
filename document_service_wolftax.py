# -*- coding: utf-8 -*-
import asyncio
from pathlib import Path
from typing import List
from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from copy import deepcopy
import re
import config_wolftax as config
from models_wolftax import WolftaxOfertaRequest


class WolftaxDocumentService:
    """Serwis do generowania dokumentów DOCX z multi-file szablonów WolfTax - używa docxcompose dla 1:1 jakości"""

    def __init__(self):
        self.templates_dir = config.TEMPLATES_DIR
        self.produkty_dir = config.PRODUKTY_DIR
        self.wolftax_files = config.WOLFTAX_FILES
        self.injection_after = config.INJECTION_AFTER_FILE

    async def generate_offer(self, request: WolftaxOfertaRequest, output_path: Path) -> Path:
        """
        Generuje ofertę DOCX na podstawie szablonów WolfTax i danych z request
        Używa docxcompose dla perfekcyjnego łączenia 1:1

        Args:
            request: Dane do wypełnienia oferty
            output_path: Ścieżka do zapisu wygenerowanego pliku

        Returns:
            Path: Ścieżka do wygenerowanego pliku DOCX
        """
        print(f"\n{'='*60}")
        print(f"🔧 WOLFTAX DOCUMENT SERVICE - Profesjonalne łączenie 1:1")
        print(f"{'='*60}\n")

        # Krok 1: Wczytaj i przygotuj pierwszy dokument (base)
        first_file = self.wolftax_files[0]["file"]
        first_path = self.templates_dir / first_file

        print(f"📄 [1/3] Ładowanie dokumentu bazowego: {first_file}")
        base_doc = Document(str(first_path))

        # Zamień placeholdery w pierwszym dokumencie
        await self._replace_placeholders(base_doc, request)

        # Utwórz Composer z dokumentem bazowym
        composer = Composer(base_doc)

        # Krok 2: Dodaj kolejne pliki szablonu
        print(f"\n📚 [2/3] Łączenie plików szablonu...")
        injection_index = self._get_injection_index()

        for i, file_info in enumerate(self.wolftax_files[1:], 1):
            file_name = file_info["file"]
            file_path = self.templates_dir / file_name

            print(f"  ➕ Dodawanie: {file_name}")

            # Wczytaj dokument
            temp_doc = Document(str(file_path))

            # Zamień placeholdery
            await self._replace_placeholders(temp_doc, request)

            # Dodaj page break PRZED nowym dokumentem
            self._add_page_break_to_composer(composer)

            # Dodaj do composer - zachowuje 100% formatowania!
            composer.append(temp_doc)

            # Sprawdź czy to jest punkt wstrzyknięcia produktów
            if file_name == self.injection_after and request.produkty:
                print(f"\n  💉 [INJECTION POINT] Wstrzykiwanie produktów...")
                await self._inject_products_with_composer(composer, request.produkty)

        # Krok 3: Zapisz wynikowy dokument
        print(f"\n💾 [3/3] Zapisywanie dokumentu...")
        composer.save(str(output_path))

        print(f"✅ Dokument zapisany: {output_path}")
        print(f"{'='*60}\n")

        return output_path

    def _get_injection_index(self) -> int:
        """Zwraca indeks pliku po którym następuje wstrzyknięcie"""
        for i, file_info in enumerate(self.wolftax_files):
            if file_info["file"] == self.injection_after:
                return i
        return 3

    async def _replace_placeholders(self, doc: Document, request: WolftaxOfertaRequest):
        """Zamienia placeholdery w dokumencie na wartości z request"""

        # Mapa placeholderów -> wartości (dopasowane do szablonu WolfTax)
        replacements = {
            "{{NazwaFirmyKlienta}}": request.nazwa_firmy_klienta,
            "{{Sygnatura-sprawy}}": request.sygnatura_sprawy or "",
            "{{Temat}}": request.temat,
            "{{Termin}}": request.termin or "",
            "{{waznosc-oferty}}": request.waznosc_oferty,
            "{{Wynagrodzenie}}": f"{request.wynagrodzenie:.2f} PLN" if request.wynagrodzenie else "",
            "{{Szacowanyczaspracy}}": str(request.szacowany_czas_pracy) if request.szacowany_czas_pracy else "",
        }

        # Zamień w paragrafach
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    await self._replace_in_paragraph(paragraph, key, value)

        # Zamień w tabelach
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                await self._replace_in_paragraph(paragraph, key, value)

    async def _replace_in_paragraph(self, paragraph, search: str, replace: str):
        """
        Zamienia tekst w paragrafie zachowując formatowanie
        Obsługuje przypadki gdy placeholder jest rozdzielony między różne runs
        """
        if search not in paragraph.text:
            return

        # Strategia: Znajdź pełny tekst, podmień, a potem przepisz do runs
        full_text = paragraph.text

        if search in full_text:
            new_text = full_text.replace(search, replace)

            # Usuń wszystkie runs
            for run in paragraph.runs:
                run.text = ""

            # Dodaj nowy tekst do pierwszego run (zachowuje formatowanie pierwszego run)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

    def _add_page_break_to_composer(self, composer: Composer):
        """
        Dodaje page break do dokumentu głównego w composer
        Każdy nowy plik/produkt zaczyna się od nowej strony
        """
        # Pobierz dokument z composer
        doc = composer.doc
        # Dodaj nowy paragraf z page break
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        # Dodaj page break używając WD_BREAK.PAGE
        run.add_break(WD_BREAK.PAGE)

    async def _inject_products_with_composer(self, composer: Composer, product_files: List[str]):
        """
        Wstrzykuje produkty jako osobne dokumenty używając Composer
        GWARANTUJE 1:1 jakość bez rozjeżdżania się elementów

        Args:
            composer: Composer z dokumentem głównym
            product_files: Lista nazw plików produktów
        """
        for product_file in product_files:
            product_path = self.produkty_dir / product_file

            if not product_path.exists():
                print(f"  ⚠️  Produkt {product_file} nie istnieje, pomijam")
                continue

            # Wczytaj dokument produktu
            product_doc = Document(str(product_path))
            print(f"    ✓ Produkt: {product_file}")

            # Dodaj page break PRZED produktem
            self._add_page_break_to_composer(composer)

            # Dodaj produkt używając Composer - 100% zachowanie formatowania!
            composer.append(product_doc)
